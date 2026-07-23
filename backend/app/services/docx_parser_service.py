"""DOCX text extraction using python-docx."""
from __future__ import annotations

import io

import docx

from app.core.logging import get_logger
from app.services.text_cleaner_service import TextCleanerService

logger = get_logger(__name__)


class DocxParserService:
    def __init__(self) -> None:
        self.cleaner = TextCleanerService()

    def extract(self, content: bytes) -> str:
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001
            raise ValueError(f"Could not read DOCX: {exc}") from exc

        parts: list[str] = [p.text for p in document.paragraphs]

        # Include table cell text too.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append(" | ".join(c for c in cells if c))

        return self.cleaner.clean("\n".join(parts))
